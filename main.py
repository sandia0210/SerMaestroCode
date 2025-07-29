import os
import re
import io
import docx
import pandas as pd
import numpy as np
from collections import Counter, defaultdict
#
# Procesamiento de texto
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer

# Google Drive API
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2 import service_account

# Topic Modelling y extracción de palabras clave
import gensim
from gensim import corpora
from gensim.models import LdaModel
from gensim.models import CoherenceModel
from keybert import KeyBERT

# Descargar recursos de NLTK si no están presentes
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
    
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab')
    
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')
    
try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('wordnet')


class GoogleDriveTopicModelling:
    def __init__(self, language='spanish'):
        """
        Inicializa el extractor de palabras clave con Google Drive integration
        
        Args:
            language (str): Idioma para stopwords ('spanish' o 'english')
        """
        self.language = language
        # Configurar stopwords según el idioma
        self.stop_words = set(stopwords.words(language if language != 'spanish' else 'spanish'))
        self.lemmatizer = WordNetLemmatizer()
        # Inicializar KeyBERT
        self.keybert_model = KeyBERT()
        
        # Google Drive API setup
        self.SCOPES = ['https://www.googleapis.com/auth/drive']
        self.service = None
        
    def authenticate_google_drive(self, credentials_file='credentials.json', token_file='token.json'):
        """
        Autentica con Google Drive API con manejo mejorado de errores
        
        Args:
            credentials_file (str): Ruta al archivo credentials.json descargado de Google Cloud Console
            token_file (str): Ruta donde se guardará el token de acceso
        """
        creds = None
        
        # Si hay un token existente y está dando problemas, eliminarlo
        if os.path.exists(token_file):
            try:
                creds = Credentials.from_authorized_user_file(token_file, self.SCOPES)
                # Verificar si las credenciales son válidas
                if creds and creds.expired and creds.refresh_token:
                    print("Token expirado, intentando renovar...")
                    creds.refresh(Request())
                elif not creds or not creds.valid:
                    print("Token inválido, eliminando y solicitando nueva autenticación...")
                    os.remove(token_file)
                    creds = None
            except Exception as e:
                print(f"Error con token existente: {e}")
                print("Eliminando token corrupto y solicitando nueva autenticación...")
                os.remove(token_file)
                creds = None
        
        # Si no hay credenciales válidas disponibles, permite al usuario autenticarse.
        if not creds or not creds.valid:
            if not os.path.exists(credentials_file):
                raise FileNotFoundError(
                    f"No se encontró el archivo {credentials_file}. "
                    "Descarga las credenciales desde Google Cloud Console."
                )
            
            print("Iniciando proceso de autenticación...")
            flow = InstalledAppFlow.from_client_secrets_file(credentials_file, self.SCOPES)
            creds = flow.run_local_server(port=0)
            
            # Guardar las credenciales para la próxima ejecución
            with open(token_file, 'w') as token:
                token.write(creds.to_json())
            print("Nuevas credenciales guardadas.")
        
        self.service = build('drive', 'v3', credentials=creds)
        print("Autenticación con Google Drive exitosa!")
    
    def reset_authentication(self, token_file='token.json'):
        """
        Elimina el token existente para forzar una nueva autenticación
        """
        if os.path.exists(token_file):
            os.remove(token_file)
            print(f"Token {token_file} eliminado. Deberás autenticarte nuevamente.")
        else:
            print("No hay token para eliminar.")

    def find_diplomado_folders(self, parent_folder_id):
        """
        Encuentra todas las carpetas que contengan "DIPLOMADO" en su nombre
        
        Args:
            parent_folder_id (str): ID de la carpeta padre
            
        Returns:
            list: Lista de carpetas de diplomados encontradas
        """
        try:
            query = f"'{parent_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and name contains 'DIPLOMADO'"
            results = self.service.files().list(
                q=query, 
                fields="files(id, name)",
                pageSize=1000
            ).execute()
            
            diplomado_folders = results.get('files', [])
            
            # Filtrar carpetas que coincidan con el patrón "#. DIPLOMADO"
            filtered_folders = []
            for folder in diplomado_folders:
                # Buscar patrón: número seguido de punto y DIPLOMADO
                if 'DIPLOMADO' in folder['name'].upper():
                    filtered_folders.append(folder)
            
            print(f"Encontradas {len(filtered_folders)} carpetas de diplomados:")
            for folder in filtered_folders:
                print(f"  - {folder['name']}")
            
            return filtered_folders
            
        except Exception as e:
            print(f"Error al buscar carpetas de diplomados: {e}")
            return []

    def navigate_to_modulo_iv(self, diplomado_folder_id):
        """
        Navega desde la carpeta del diplomado hasta MÓDULO IV
        Ruta: DIPLOMADO -> 6. EVIDENCIA DE TRABAJOS -> MÓDULO IV
        
        Args:
            diplomado_folder_id (str): ID de la carpeta del diplomado
            
        Returns:
            str: ID de la carpeta MÓDULO IV o None si no se encuentra
        """
        try:
            # Paso 1: Buscar carpeta "6. EVIDENCIA DE TRABAJOS"
            query = f"'{diplomado_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
            results = self.service.files().list(q=query, fields="files(id, name)").execute()
            folders = results.get('files', [])
            
            evidencia_folder = None
            for folder in folders:
                if 'EVIDENCIA' in folder['name'].upper() and 'TRABAJOS' in folder['name'].upper():
                    evidencia_folder = folder
                    break
            
            if not evidencia_folder:
                print("No se encontró la carpeta de EVIDENCIA DE TRABAJOS")
                return None
            
            # Paso 2: Buscar carpeta "MÓDULO IV" dentro de EVIDENCIA DE TRABAJOS
            query = f"'{evidencia_folder['id']}' in parents and mimeType='application/vnd.google-apps.folder'"
            results = self.service.files().list(q=query, fields="files(id, name)").execute()
            folders = results.get('files', [])
            
            modulo_iv_folder = None
            for folder in folders:
                if 'MÓDULO' in folder['name'].upper() and 'IV' in folder['name'].upper():
                    modulo_iv_folder = folder
                    break
            
            if not modulo_iv_folder:
                print("No se encontró la carpeta MÓDULO IV")
                return None
            
            return modulo_iv_folder['id']
            
        except Exception as e:
            print(f"Error al navegar a MÓDULO IV: {e}")
            return None

    def get_folders_by_pattern_improved(self, parent_folder_id):
        """
        Encuentra TODOS los grupos en la carpeta MÓDULO IV
        """
        try:
            query = f"'{parent_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
            
            results = self.service.files().list(
                q=query, 
                fields="files(id, name)",
                pageSize=1000
            ).execute()
            
            folders = results.get('files', [])
            
            # Filtrar y procesar carpetas de grupo
            group_folders = {}
            
            for folder in folders:
                folder_name = folder['name']
                
                # Solo procesar carpetas que contengan "grupo" (case insensitive)
                if 'grupo' not in folder_name.lower():
                    continue
                
                # Patrones más flexibles para extraer números
                patterns = [
                    r'[Gg]rupo\s*0*(\d+)',         # "Grupo 01", "grupo 1", "GRUPO 001"
                    r'[Gg]rupo\s*(\d+)',           # "Grupo1", "grupo23"
                    r'(\d+).*[Gg]rupo',            # "01 Grupo", "1-Grupo"
                    r'(\d+)',                      # Cualquier número en el nombre
                ]
                
                group_number = None
                for pattern in patterns:
                    match = re.search(pattern, folder_name)
                    if match:
                        # Remover ceros a la izquierda pero mantener al menos un dígito
                        group_number = match.group(1).lstrip('0') or '0'
                        break
                
                if group_number:
                    # Usar el número como clave para evitar duplicados
                    if group_number not in group_folders:
                        group_folders[group_number] = folder
                    else:
                        # Si hay múltiples carpetas con el mismo número, elegir la más "estándar"
                        current_name = group_folders[group_number]['name']
                        new_name = folder['name']
                        
                        # Preferir nombres más estándar (con "Grupo" al inicio)
                        if (new_name.lower().startswith('grupo') and 
                            not current_name.lower().startswith('grupo')):
                            group_folders[group_number] = folder
            
            return list(group_folders.values()), group_folders
            
        except Exception as e:
            print(f"Error al obtener carpetas mejorado: {e}")
            return [], {}

    def find_sistematizacion_file(self, folder_id):
        """
        Busca un archivo .docx que contenga "SISTEMATIZACION" en su nombre dentro de la carpeta
        
        Args:
            folder_id (str): ID de la carpeta donde buscar
            
        Returns:
            dict: Información del archivo encontrado o None
        """
        try:
            query = f"'{folder_id}' in parents and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'"
            results = self.service.files().list(q=query, fields="files(id, name)").execute()
            files = results.get('files', [])
            
            for file in files:
                if ('SISTEMATIZACION' in file['name'].upper()) or (('SISTEMATIZACIÓN' in file['name'].upper())):
                    # Obtener información adicional del archivo incluyendo webViewLink
                    file_details = self.service.files().get(
                        fileId=file['id'], 
                        fields="id, name, webViewLink"
                    ).execute()
                    return file_details
            
            return None
            
        except Exception as e:
            print(f"Error al buscar archivo: {e}")
            return None
    
    def download_file_content(self, file_id):
        """
        Descarga el contenido de un archivo de Google Drive
        
        Args:
            file_id (str): ID del archivo a descargar
            
        Returns:
            bytes: Contenido del archivo
        """
        try:
            request = self.service.files().get_media(fileId=file_id)
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            return file_content.getvalue()
            
        except Exception as e:
            print(f"Error al descargar archivo: {e}")
            return None

    def extraer_titulo_proyecto_from_bytes(self, file_bytes, filename):
        """
        Extrae el título del proyecto desde la primera tabla del documento
        Busca la celda que está a la derecha de "TÍTULO" o "TITULO"
        
        Args:
            file_bytes (bytes): Contenido del archivo DOCX en bytes
            filename (str): Nombre del archivo para logging
        
        Returns:
            str: Título del proyecto extraído
        """
        try:
            # Crear un objeto BytesIO desde los bytes
            file_stream = io.BytesIO(file_bytes)
            
            # Cargar el documento desde el stream
            doc = docx.Document(file_stream)
            
            # Buscar en todas las tablas del documento
            for tabla in doc.tables:
                # Recorrer todas las filas de la tabla
                for fila in tabla.rows:
                    # Recorrer todas las celdas de la fila
                    for i, celda in enumerate(fila.cells):
                        celda_text = celda.text.strip().upper()
                        
                        # Verificar si la celda contiene "TÍTULO" o "TITULO"
                        if 'TÍTULO' in celda_text or 'TITULO' in celda_text:
                            # Intentar obtener la celda siguiente (a la derecha)
                            if i + 1 < len(fila.cells):
                                titulo = fila.cells[i + 1].text.strip()
                                if titulo and len(titulo) > 5:  # Verificar que no esté vacío
                                    return titulo
                            
                            # Si no hay celda a la derecha, intentar en la misma celda después del texto "TÍTULO"
                            if ':' in celda.text:
                                partes = celda.text.split(':', 1)
                                if len(partes) > 1:
                                    titulo = partes[1].strip()
                                    if titulo and len(titulo) > 5:
                                        return titulo
            
            # Si no se encuentra en tablas, buscar en párrafos
            for para in doc.paragraphs:
                text = para.text.strip()
                if ('TÍTULO' in text.upper() or 'TITULO' in text.upper()) and ':' in text:
                    partes = text.split(':', 1)
                    if len(partes) > 1:
                        titulo = partes[1].strip()
                        if titulo and len(titulo) > 5:
                            return titulo
            
            print(f"No se pudo extraer el título del proyecto de {filename}")
            return "TÍTULO NO ENCONTRADO"
            
        except Exception as e:
            print(f"Error al extraer título de {filename}: {e}")
            return "ERROR AL EXTRAER TÍTULO"

    def extraer_resumen_ejecutivo_from_bytes(self, file_bytes, filename):
        """
        Extrae el contenido de la primera celda de la tabla que contiene el Resumen Ejecutivo
        desde bytes del archivo.
        
        Args:
            file_bytes (bytes): Contenido del archivo DOCX en bytes
            filename (str): Nombre del archivo para logging
        
        Returns:
            str: Texto extraído de la primera celda de la tabla del Resumen Ejecutivo
        """
        try:
            # Crear un objeto BytesIO desde los bytes
            file_stream = io.BytesIO(file_bytes)
            
            # Cargar el documento desde el stream
            doc = docx.Document(file_stream)
            
            # Variable para indicar que hemos encontrado la sección
            seccion_encontrada = False
            
            # Buscar el párrafo que contiene "1. Resumen ejecutivo"
            for i, para in enumerate(doc.paragraphs):
                if "Resumen ejecutivo" in para.text:
                    seccion_encontrada = True
                    break
            
            if not seccion_encontrada:
                print(f"No se encontró la sección de Resumen Ejecutivo en {filename}")
                # Si no hay resumen ejecutivo, obtener todo el texto del documento
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():  # Solo añadir párrafos no vacíos
                        full_text.append(para.text)
                return ' '.join(full_text)
            
            # Buscar la tabla que sigue después del título
            for tabla in doc.tables:
                # Verificamos si estamos en la tabla correcta (la que contiene el resumen)
                texto_primera_celda = ""
                
                # Recorrer las filas de la tabla
                for fila in tabla.rows:
                    # Verificar si hay celdas
                    if fila.cells and len(fila.cells) > 0:
                        celda = fila.cells[0]  # Primera celda
                        
                        # Extraer texto de la primera celda
                        texto_primera_celda = celda.text.strip()
                        
                        # Verificar si este texto contiene contenido del resumen ejecutivo
                        if len(texto_primera_celda) > 50:  # Bajamos el umbral para ser más flexibles
                            return texto_primera_celda
            
            # Si no se encontró en tablas, intentar extraer texto de párrafos después de "Resumen ejecutivo"
            full_text = []
            found_section = False
            next_section = False
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not found_section and "Resumen ejecutivo" in text:
                    found_section = True
                    continue
                elif found_section and text and not next_section:
                    # Verificar si llegamos a la siguiente sección (usualmente comienza con un número)
                    if re.match(r'^\d+\.', text) and "Resumen" not in text:
                        next_section = True
                    else:
                        full_text.append(text)
                elif next_section:
                    break
            
            if full_text:
                return ' '.join(full_text)
            
            # Si todo falla, extraer todo el contenido del documento
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():  # Solo añadir párrafos no vacíos
                    full_text.append(para.text)
            
            print(f"No se pudo extraer el resumen específico de {filename}, usando texto completo")
            return ' '.join(full_text)
            
        except Exception as e:
            print(f"Error al procesar {filename}: {e}")
            return ""

    def preprocess_text(self, text):
        """
        Preprocesa el texto para análisis
        
        Args:
            text (str): Texto a preprocesar
            
        Returns:
            list: Lista de tokens procesados
        """
        # Verificar si text es None o vacío
        if not text:
            return []
            
        # Convertir a minúsculas y eliminar caracteres especiales
        text = text.lower()
        text = re.sub(r'[^\w\s]', '', text)
        
        # Tokenizar el texto
        tokens = word_tokenize(text)
        
        # Eliminar stopwords y palabras cortas
        tokens = [token for token in tokens if token not in self.stop_words and len(token) > 2]
        
        # Lematizar las palabras
        if self.language == 'english':
            tokens = [self.lemmatizer.lemmatize(token) for token in tokens]
        
        return tokens
    
    def extract_keywords_keybert(self, text, top_n=10):
        """
        Extrae palabras clave usando KeyBERT
        
        Args:
            text (str): Texto del documento
            top_n (int): Número de palabras clave a extraer
            
        Returns:
            list: Lista de tuplas (palabra_clave, puntuación)
        """
        try:
            if not text or len(text.strip()) < 100:  # Verificar que hay suficiente texto
                print("Texto insuficiente para extraer palabras clave")
                return []
                
            keywords = self.keybert_model.extract_keywords(
                text, 
                keyphrase_ngram_range=(1, 2), 
                stop_words=list(self.stop_words),
                top_n=top_n
            )
            return keywords
        except Exception as e:
            print(f"Error al extraer palabras clave: {e}")
            return []

    def process_single_diplomado(self, diplomado_folder):
        """
        Procesa un diplomado individual
        
        Args:
            diplomado_folder (dict): Información de la carpeta del diplomado
            
        Returns:
            list: Lista de registros para este diplomado
        """
        diplomado_name = diplomado_folder['name']
        print(f"\n=== PROCESANDO DIPLOMADO: {diplomado_name} ===")
        
        # Navegar hasta MÓDULO IV
        modulo_iv_id = self.navigate_to_modulo_iv(diplomado_folder['id'])
        
        if not modulo_iv_id:
            print(f"No se pudo acceder a MÓDULO IV en {diplomado_name}")
            return []
        
        # Obtener todas las carpetas de grupo
        all_group_folders, group_folders_dict = self.get_folders_by_pattern_improved(modulo_iv_id)
        
        if not group_folders_dict:
            print(f"No se encontraron grupos en {diplomado_name}")
            return []
        
        # Lista para almacenar registros de este diplomado
        diplomado_records = []
        
        # Procesar cada grupo
        group_numbers = sorted(group_folders_dict.keys(), key=int)
        print(f"Procesando {len(group_numbers)} grupos en {diplomado_name}")
        
        for group_num in group_numbers:
            folder = group_folders_dict[group_num]
            print(f"  Procesando Grupo {group_num}: {folder['name']}")
            
            # Buscar archivo de sistematización
            sistematizacion_file = self.find_sistematizacion_file(folder['id'])
            
            if not sistematizacion_file:
                print(f"    ❌ No se encontró archivo de sistematización")
                continue
            
            print(f"    ✅ Archivo encontrado: {sistematizacion_file['name']}")
            
            # Descargar contenido
            file_content = self.download_file_content(sistematizacion_file['id'])
            
            if not file_content:
                print(f"    ❌ Error al descargar archivo")
                continue
            
            # Extraer título del proyecto
            titulo_proyecto = self.extraer_titulo_proyecto_from_bytes(file_content, sistematizacion_file['name'])
            
            # Extraer texto del resumen ejecutivo
            text = self.extraer_resumen_ejecutivo_from_bytes(file_content, sistematizacion_file['name'])
            
            if not text or len(text.strip()) < 50:
                print(f"    ❌ Texto insuficiente para análisis")
                continue
            
            # Extraer keywords
            keywords_with_scores = self.extract_keywords_keybert(text, top_n=5)
            
            if not keywords_with_scores:
                print(f"    ❌ No se pudieron extraer keywords")
                continue
            
            # Crear registro
            keywords_list = [keyword for keyword, score in keywords_with_scores]
            
            # Construir enlace de descarga
            download_link = f"https://docs.google.com/document/d/{sistematizacion_file['id']}/export?format=docx"

            record = {
                'Diplomado': diplomado_name,
                'Nombre de documento': sistematizacion_file['name'],
                'Título del proyecto': titulo_proyecto,
                'Enlace de descarga': download_link
            }
            
            # Agregar keywords (máximo 5)
            for i in range(5):
                key_name = f'keyword {i+1}'
                if i < len(keywords_list):
                    record[key_name] = keywords_list[i]
                else:
                    record[key_name] = ""
            
            diplomado_records.append(record)
            print(f"    ✅ Registro creado exitosamente")
        
        return diplomado_records

    def process_all_diplomados(self, parent_folder_id, top_keywords=5):
        """
        Procesa todos los diplomados encontrados en la carpeta padre
        
        Args:
            parent_folder_id (str): ID de la carpeta padre que contiene los diplomados
            top_keywords (int): Número de palabras clave por documento (máximo 5)
            
        Returns:
            pd.DataFrame: DataFrame con todos los resultados
        """
        if not self.service:
            raise Exception("Primero debes autenticarte con Google Drive usando authenticate_google_drive()")
        
        # Encontrar todas las carpetas de diplomados
        diplomado_folders = self.find_diplomado_folders(parent_folder_id)
        
        if not diplomado_folders:
            print("No se encontraron carpetas de diplomados!")
            return pd.DataFrame()
        
        # Lista para almacenar todos los registros
        all_records = []
        
        # Procesar cada diplomado
        for diplomado_folder in diplomado_folders:
            diplomado_records = self.process_single_diplomado(diplomado_folder)
            all_records.extend(diplomado_records)
        
        # Crear DataFrame
        if all_records:
            df = pd.DataFrame(all_records)
            
            # Reordenar columnas
            columns_order = ['Diplomado', 'Nombre de documento', 'Título del proyecto', 'Enlace de descarga'] + [f'keyword {i+1}' for i in range(5)]
            df = df[columns_order]
            
            print(f"\n=== RESUMEN FINAL ===")
            print(f"Total de diplomados procesados: {len(diplomado_folders)}")
            print(f"Total de proyectos procesados: {len(df)}")
            print(f"Proyectos por diplomado:")
            diplomado_counts = df['Diplomado'].value_counts()
            for diplomado, count in diplomado_counts.items():
                print(f"  - {diplomado}: {count} proyectos")
            
            return df
        else:
            print("No se procesaron proyectos exitosamente")
            return pd.DataFrame()
        
    def upload_excel_to_drive(self, excel_filename, parent_folder_id, drive_filename=None):
        """
        Sube un archivo Excel a Google Drive, sobreescribiendo si ya existe
        
        Args:
            excel_filename (str): Ruta local del archivo Excel
            parent_folder_id (str): ID de la carpeta destino en Drive
            drive_filename (str): Nombre del archivo en Drive (opcional)
        
        Returns:
            str: ID del archivo subido o None si hay error
        """
        try:
            from googleapiclient.http import MediaFileUpload
            
            if not drive_filename:
                drive_filename = excel_filename
            
            # Buscar si ya existe un archivo con el mismo nombre
            query = f"name='{drive_filename}' and '{parent_folder_id}' in parents and trashed=false"
            results = self.service.files().list(q=query, fields="files(id, name)").execute()
            existing_files = results.get('files', [])
            
            media = MediaFileUpload(excel_filename, 
                                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            
            if existing_files:
                # Si existe, actualizar el archivo existente (mantiene el mismo ID)
                file_id = existing_files[0]['id']
                print(f"📝 Archivo existente encontrado, actualizando ID: {file_id}")
                
                file = self.service.files().update(
                    fileId=file_id,
                    media_body=media,
                    fields='id'
                ).execute()
                
                print(f"✅ Archivo actualizado exitosamente, mismo ID: {file.get('id')}")
                
            else:
                # Si no existe, crear nuevo archivo
                print(f"📁 Creando nuevo archivo: {drive_filename}")
                
                file_metadata = {
                    'name': drive_filename,
                    'parents': [parent_folder_id]
                }
                
                file = self.service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields='id'
                ).execute()
                
                print(f"✅ Nuevo archivo creado con ID: {file.get('id')}")
            
            return file.get('id')
            
        except Exception as e:
            print(f"❌ Error al subir/actualizar archivo: {e}")
            return None


def reset_auth_and_run_multi():
    """
    Función para resetear autenticación y ejecutar el script para múltiples diplomados
    """
    print("=== RESETEANDO AUTENTICACIÓN ===")
    topic_model = GoogleDriveTopicModelling(language='spanish')
    topic_model.reset_authentication() # Elimina token.json
    
    try:
        # Forzar nueva autenticación
        topic_model.authenticate_google_drive()
        
        # ID de la carpeta padre (que contiene los diplomados)
        parent_folder_id = "1-_W-Esk4lzkztPSeZpqO4Gq3ao1P9XKo"  # Cambiar por tu ID
        
        # Procesar todos los diplomados
        result_df = topic_model.process_all_diplomados(parent_folder_id, top_keywords=5)
        
        return result_df
        
    except Exception as e:
        print(f"Error: {e}")
        return pd.DataFrame()


def main():
    """
    Función principal que ejecuta el procesamiento automáticamente para múltiples diplomados
    """
    # Inicializar el modelo
    topic_model = GoogleDriveTopicModelling(language='spanish')
    
    try:
        # Autenticar con Google Drive
        topic_model.authenticate_google_drive()
        
        # ID de la carpeta padre (que contiene los diplomados)
        parent_folder_id = "1-_W-Esk4lzkztPSeZpqO4Gq3ao1P9XKo"  # Cambiar por tu ID
        
        # Procesar todos los diplomados
        result_df = topic_model.process_all_diplomados(parent_folder_id, top_keywords=5)
        
        return result_df
        
    except Exception as e:
        print(f"Error en main: {e}")
        if "invalid_grant" in str(e).lower():
            print("\n=== SOLUCIÓN SUGERIDA ===")
            print("El error 'invalid_grant' indica que tu token ha expirado.")
            print("Ejecuta: reset_auth_and_run_multi() para solucionarlo.")
        return pd.DataFrame()


if __name__ == "__main__":
    # Ejecutar procesamiento
    result_df = reset_auth_and_run_multi()
    
    if not result_df.empty:
        print(f"\n=== GUARDANDO RESULTADOS ===")
        
        # Mostrar preview del DataFrame
        print("\nPreview del DataFrame:")
        print(result_df.head())
        
        # Guardar en Excel
        output_filename = "output_multi_diplomados_completo.xlsx"
        result_df.to_excel(output_filename, index=False)
        print(f"\nResultados guardados en '{output_filename}'")
        
        #Meter el excel a la carpeta
        topic_model = GoogleDriveTopicModelling(language='spanish')
        topic_model.authenticate_google_drive()
        parent_folder_id = "1-_W-Esk4lzkztPSeZpqO4Gq3ao1P9XKo"
        
        uploaded_file_id = topic_model.upload_excel_to_drive(
            output_filename, 
            parent_folder_id, 
            f"Resultados_Keywords.xlsx"
        )
        
        if uploaded_file_id:
            print(f"✅ Archivo también guardado en Google Drive")


        # Mostrar estadísticas finales
        print(f"\nEstadísticas finales:")
        print(f"- Total de proyectos: {len(result_df)}")
        print(f"- Diplomados únicos: {result_df['Diplomado'].nunique()}")
        
    else:
        print("No se procesaron documentos exitosamente.")